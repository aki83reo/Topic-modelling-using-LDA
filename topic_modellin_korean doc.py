'''
  libraries
'''

import spacy
import nltk
import pandas as pd
spacy.load('en_core_web_sm')
from spacy.lang.en import English
parser = English()
from docx import Document
from  nltk import tokenize
from nltk.corpus import wordnet as wn
nltk.download('wordnet')



def topic_modelling_doc_files(file_path):
    '''

    lemetizing words  :  bring words to  its original form

    '''
    def get_lemma(word):
        lemma = wn.morphy(word)
        if lemma is None:
            return word
        else:
            return lemma

    '''
    stopwords removel :  remove  stopwords  from  the corpus 
    '''

    #nltk.download('stopwords')
    en_stop = set(nltk.corpus.stopwords.words('english'))

    '''
     Below function  is to  apply  ,  tokenizing , applying lemmatizing  ,  removing stopwords 
    '''

    def prepare_text_for_lda(text):
        tokens = nltk.word_tokenize(text)
        tokens = [token for token in tokens if len(token) > 4]
        tokens = [token for token in tokens if token not in en_stop]
        tokens = [get_lemma(token) for token in tokens]
        return tokens


    '''
      Call  any  docx  file  
    '''

    list_ver=[]
    #document = Document("D:/preprocessed.docx")
    document = Document(file_path)

    for para in document.paragraphs:
        list_ver.append(para.text)

    data_into_dataframe_val=pd.DataFrame({"data_list":list_ver})  ##  list 2 df


    new_text_for_lda=[]
    for  i  in range(data_into_dataframe_val.__len__()):
        new_text_for_lda.append(prepare_text_for_lda(data_into_dataframe_val.data_list.iloc[i]))   ##  tokenize dataframe into list

    ## LDA using  gensim

    '''
    First, we are creating a dictionary from the data, then convert to bag-of-words corpus and save the dictionary and corpus for future use.
    '''

    from gensim import corpora
    dictionary = corpora.Dictionary(new_text_for_lda)
    corpus = [dictionary.doc2bow(text) for text in new_text_for_lda]
    import pickle
    pickle.dump(corpus, open('corpus.pkl', 'wb'))
    dictionary.save('dictionary.gensim')

    '''
    We are asking LDA to find 10 topics in the data:
    '''

    import gensim
    NUM_TOPICS = 10
    ldamodel = gensim.models.ldamodel.LdaModel(corpus, num_topics = NUM_TOPICS, id2word=dictionary, passes=15)
    ldamodel.save('model5.gensim')


    topic_list=[]
    topics = ldamodel.print_topics(num_words=4)
    for topic in topics:
        print(topic)
        topic_list.append(topic)

    return topic_list



    '''
    import pickle
    dictionary = gensim.corpora.Dictionary.load('dictionary.gensim')
    corpus = pickle.load(open('corpus.pkl', 'rb'))
    lda = gensim.models.ldamodel.LdaModel.load('model5.gensim')
    
    import pyLDAvis.gensim
    import IPython
    lda_display = pyLDAvis.gensim.prepare(lda, corpus, dictionary, sort_topics=False)
    pyLDAvis.display(lda_display)
    
    '''

    '''
    
   
